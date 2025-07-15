"""
Document processor for handling multiple file formats.
Supports: .xlsx, .doc, .docx, .pdf, images (jpg, png, etc.), .md, .txt
"""

import os
import logging
import asyncio
from typing import Dict, Any, Optional, List, Tuple
from pathlib import Path
import base64
import io

# Document processing libraries
import pandas as pd
from docx import Document
import PyPDF2
from PIL import Image

# Optional dependencies
try:
    import pytesseract
    HAS_PYTESSERACT = True
except ImportError:
    HAS_PYTESSERACT = False

try:
    from markitdown import MarkItDown
    HAS_MARKITDOWN = True
except ImportError:
    HAS_MARKITDOWN = False

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Multi-format document processor that converts various file types to text.
    """
    
    def __init__(self):
        """Initialize the document processor."""
        # Core supported extensions (always available)
        self.supported_extensions = {
            '.xlsx', '.xls',  # Excel
            '.docx', '.doc',  # Word
            '.pdf',           # PDF
            '.md', '.markdown', '.txt'  # Text files
        }
        
        # Add image support if pytesseract is available
        if HAS_PYTESSERACT:
            self.supported_extensions.update({
                '.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.gif'  # Images
            })
        
        # Initialize markitdown if available
        self.markitdown_processor = MarkItDown() if HAS_MARKITDOWN else None
    
    def is_supported(self, file_path: str) -> bool:
        """
        Check if a file is supported by this processor.
        
        Args:
            file_path: Path to the file
            
        Returns:
            True if supported, False otherwise
        """
        return Path(file_path).suffix.lower() in self.supported_extensions
    
    async def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        Process a document and extract text content.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Dictionary containing:
            - content: Extracted text content
            - metadata: Document metadata
            - file_type: Type of file processed
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = Path(file_path).suffix.lower()
        
        # Try markitdown first if available (it handles many formats well)
        if self.markitdown_processor and file_extension in {'.xlsx', '.xls', '.docx', '.doc', '.pdf'}:
            try:
                return await self._process_with_markitdown(file_path)
            except Exception as e:
                logger.warning(f"Markitdown processing failed for {file_path}: {e}")
                # Fall back to specific processors
        
        # Use specific processors based on file type
        if file_extension in {'.xlsx', '.xls'}:
            return await self._process_excel(file_path)
        elif file_extension in {'.docx', '.doc'}:
            return await self._process_word(file_path)
        elif file_extension == '.pdf':
            return await self._process_pdf(file_path)
        elif file_extension in {'.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.gif'}:
            if HAS_PYTESSERACT:
                return await self._process_image(file_path)
            else:
                raise ValueError(f"Image processing not available - pytesseract not installed")
        elif file_extension in {'.md', '.markdown', '.txt'}:
            return await self._process_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    async def _process_with_markitdown(self, file_path: str) -> Dict[str, Any]:
        """
        Process document using markitdown library.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Document processing result
        """
        def _process_sync():
            result = self.markitdown_processor.convert(file_path)
            return result.text_content
        
        # Run in thread pool to avoid blocking
        content = await asyncio.get_event_loop().run_in_executor(None, _process_sync)
        
        return {
            'content': content,
            'metadata': {
                'processor': 'markitdown',
                'file_path': file_path,
                'file_size': os.path.getsize(file_path),
                'file_type': Path(file_path).suffix.lower()
            },
            'file_type': Path(file_path).suffix.lower()
        }
    
    async def _process_excel(self, file_path: str) -> Dict[str, Any]:
        """
        Process Excel files (.xlsx, .xls).
        
        Args:
            file_path: Path to the Excel file
            
        Returns:
            Document processing result
        """
        def _process_sync():
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            content_parts = []
            
            for sheet_name in excel_file.sheet_names:
                try:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    
                    # Convert to markdown-like format
                    content_parts.append(f"## Sheet: {sheet_name}")
                    content_parts.append("")
                    
                    # Add table content
                    if not df.empty:
                        # Convert DataFrame to markdown table
                        markdown_table = df.to_markdown(index=False)
                        content_parts.append(markdown_table)
                    else:
                        content_parts.append("*(Empty sheet)*")
                    
                    content_parts.append("")
                    
                except Exception as e:
                    logger.warning(f"Error processing sheet {sheet_name}: {e}")
                    content_parts.append(f"## Sheet: {sheet_name} (Error: {e})")
                    content_parts.append("")
            
            return "\n".join(content_parts)
        
        # Run in thread pool to avoid blocking
        content = await asyncio.get_event_loop().run_in_executor(None, _process_sync)
        
        return {
            'content': content,
            'metadata': {
                'processor': 'pandas',
                'file_path': file_path,
                'file_size': os.path.getsize(file_path),
                'file_type': Path(file_path).suffix.lower(),
                'sheets': pd.ExcelFile(file_path).sheet_names
            },
            'file_type': Path(file_path).suffix.lower()
        }
    
    async def _process_word(self, file_path: str) -> Dict[str, Any]:
        """
        Process Word documents (.docx, .doc).
        
        Args:
            file_path: Path to the Word document
            
        Returns:
            Document processing result
        """
        def _process_sync():
            if file_path.lower().endswith('.docx'):
                # Process .docx files
                doc = Document(file_path)
                content_parts = []
                
                for paragraph in doc.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        content_parts.append(text)
                
                # Extract tables
                for table in doc.tables:
                    table_data = []
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            row_data.append(cell.text.strip())
                        table_data.append(row_data)
                    
                    if table_data:
                        # Convert to markdown table
                        content_parts.append("")
                        content_parts.append("| " + " | ".join(table_data[0]) + " |")
                        content_parts.append("| " + " | ".join(["---"] * len(table_data[0])) + " |")
                        for row in table_data[1:]:
                            content_parts.append("| " + " | ".join(row) + " |")
                        content_parts.append("")
                
                return "\n".join(content_parts)
            else:
                # For .doc files, we need a different approach
                # This is a simplified fallback
                raise ValueError("Legacy .doc format requires additional libraries (python-docx only supports .docx)")
        
        # Run in thread pool to avoid blocking
        content = await asyncio.get_event_loop().run_in_executor(None, _process_sync)
        
        return {
            'content': content,
            'metadata': {
                'processor': 'python-docx',
                'file_path': file_path,
                'file_size': os.path.getsize(file_path),
                'file_type': Path(file_path).suffix.lower()
            },
            'file_type': Path(file_path).suffix.lower()
        }
    
    async def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        Process PDF files.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Document processing result
        """
        def _process_sync():
            content_parts = []
            
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    num_pages = len(pdf_reader.pages)
                    
                    for page_num in range(num_pages):
                        page = pdf_reader.pages[page_num]
                        text = page.extract_text()
                        
                        if text.strip():
                            content_parts.append(f"## Page {page_num + 1}")
                            content_parts.append("")
                            content_parts.append(text.strip())
                            content_parts.append("")
                    
                    return "\n".join(content_parts)
                    
            except Exception as e:
                logger.error(f"Error processing PDF {file_path}: {e}")
                return f"*(Error processing PDF: {e})*"
        
        # Run in thread pool to avoid blocking
        content = await asyncio.get_event_loop().run_in_executor(None, _process_sync)
        
        return {
            'content': content,
            'metadata': {
                'processor': 'PyPDF2',
                'file_path': file_path,
                'file_size': os.path.getsize(file_path),
                'file_type': Path(file_path).suffix.lower()
            },
            'file_type': Path(file_path).suffix.lower()
        }
    
    async def _process_image(self, file_path: str) -> Dict[str, Any]:
        """
        Process image files using OCR.
        
        Args:
            file_path: Path to the image file
            
        Returns:
            Document processing result
        """
        def _process_sync():
            try:
                if not HAS_PYTESSERACT:
                    return "*(pytesseract not available - cannot process images)*"
                
                # Open and process image
                image = Image.open(file_path)
                
                # Convert to RGB if necessary
                if image.mode != 'RGB':
                    image = image.convert('RGB')
                
                # Extract text using OCR
                text = pytesseract.image_to_string(image)
                
                # Clean up the text
                lines = [line.strip() for line in text.split('\n') if line.strip()]
                
                if lines:
                    content = "\n".join(lines)
                else:
                    content = "*(No text extracted from image)*"
                
                return content
                
            except Exception as e:
                logger.error(f"Error processing image {file_path}: {e}")
                return f"*(Error processing image: {e})*"
        
        # Run in thread pool to avoid blocking
        content = await asyncio.get_event_loop().run_in_executor(None, _process_sync)
        
        return {
            'content': content,
            'metadata': {
                'processor': 'pytesseract',
                'file_path': file_path,
                'file_size': os.path.getsize(file_path),
                'file_type': Path(file_path).suffix.lower()
            },
            'file_type': Path(file_path).suffix.lower()
        }
    
    async def _process_text(self, file_path: str) -> Dict[str, Any]:
        """
        Process text files (.md, .markdown, .txt).
        
        Args:
            file_path: Path to the text file
            
        Returns:
            Document processing result
        """
        def _process_sync():
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except UnicodeDecodeError:
                # Try with different encoding
                with open(file_path, 'r', encoding='latin-1') as f:
                    return f.read()
        
        # Run in thread pool to avoid blocking
        content = await asyncio.get_event_loop().run_in_executor(None, _process_sync)
        
        return {
            'content': content,
            'metadata': {
                'processor': 'text_reader',
                'file_path': file_path,
                'file_size': os.path.getsize(file_path),
                'file_type': Path(file_path).suffix.lower()
            },
            'file_type': Path(file_path).suffix.lower()
        }
    
    def get_supported_extensions(self) -> List[str]:
        """
        Get list of supported file extensions.
        
        Returns:
            List of supported extensions
        """
        return sorted(list(self.supported_extensions))


# Factory function for creating document processor
def create_document_processor() -> DocumentProcessor:
    """
    Create a document processor instance.
    
    Returns:
        DocumentProcessor instance
    """
    return DocumentProcessor()